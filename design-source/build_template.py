"""
SAMUDRA.Jobfit DOCX template builder
------------------------------------
Builds the R-series master template that reproduces the uploaded
SAMUDRA.Jobfit PDF layout page-by-page.

Two modes:
  template -> emits docxtemplater tags ({name}, {%radarChart}, loops)
  sample   -> emits Priyanka K's data + real chart PNGs for visual QA

Run:  python3 build_template.py template out.docx
      python3 build_template.py sample  out.docx
"""
import os, sys, copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement

YELLOW   = "FFC425"
CREAM    = "FFF6E3"
CREAM2   = "FDF3D7"
SLATE    = RGBColor(0x4A, 0x65, 0x72)   # big titles
STEEL    = RGBColor(0x3D, 0x6B, 0x8E)   # section headings
BODY     = RGBColor(0x4A, 0x67, 0x85)   # body text
DARK     = RGBColor(0x33, 0x33, 0x33)
ORANGE   = RGBColor(0xF5, 0xA6, 0x23)
NAVY     = RGBColor(0x1F, 0x3F, 0x94)
BLUEBAR  = "0F9BE8"
NAVYBADGE= "1F3864"
LIGHTBLUE= "D9E9F7"

ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)

def cell_borders(cell, color="F5A623", size=8, edges=("top","left","bottom","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def fixed_widths(table, widths_in):
    """Force fixed table layout with explicit grid columns (LibreOffice honors this)."""
    from docx.shared import Inches as _In
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(sum(widths_in)*1440)))
    tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid"); tbl.insert(1, grid)
    for gc in list(grid): grid.remove(gc)
    for w in widths_in:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(w*1440)))
        grid.append(gc)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = _In(w)

def para(container, text="", size=11, color=BODY, bold=False, align=None,
         space_after=6, space_before=0, italic=False):
    p = container.add_paragraph() if hasattr(container, "add_paragraph") else container
    if text != "" or True:
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.bold = bold; r.font.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align: p.alignment = align
    return p

def brand_header(doc):
    """Yellow band with 'Powered by GPS7' + SAMUDRA wordmark, then page title."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    t.columns[0].width = Inches(4.9); t.columns[1].width = Inches(2.1)
    c0, c1 = t.rows[0].cells
    c0.width = Inches(4.9); c1.width = Inches(2.1)
    fixed_widths(t, [4.9, 2.1])
    set_cell_bg(c0, YELLOW)
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Powered by GPS7  "); r.font.bold = True; r.font.italic = True
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = c1.paragraphs[0]
    r = p2.add_run("S.A.M.U.D.R.A"); r.font.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = SLATE
    p3 = c1.add_paragraph(); r = p3.add_run("PERFORMANCE.Intelligence")
    r.font.bold = True; r.font.size = Pt(8); r.font.color.rgb = SLATE
    p = doc.add_paragraph()
    r = p.add_run("SAMUDRA"); r.font.size = Pt(24); r.font.color.rgb = SLATE
    r7 = p.add_run("7"); r7.font.size = Pt(12); r7.font.color.rgb = SLATE; r7.font.superscript = True
    r2 = p.add_run(". Jobfit"); r2.font.size = Pt(24); r2.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(4)

def section_title(doc, text):
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)
    c0, c1 = t.rows[0].cells
    c0.width = Inches(2.6); c1.width = Inches(4.4)
    fixed_widths(t, [3.0, 4.0])
    p = c0.paragraphs[0]
    r = p.add_run(text); r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = STEEL
    set_cell_bg(c1, BLUEBAR)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def yellow_rule(doc):
    t = doc.add_table(rows=1, cols=1); no_borders(t)
    fixed_widths(t, [7.0])
    set_cell_bg(t.rows[0].cells[0], YELLOW)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight"); h.set(qn("w:val"), "60"); h.set(qn("w:hRule"), "exact")
    trPr.append(h)

def trait_block(doc, title, body, size_t=12.5):
    p = doc.add_paragraph()
    r = p.add_run(title); r.font.size = Pt(size_t); r.font.bold = True; r.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(2)
    para(doc, body, size=10.5, color=BODY, space_after=8)

def image_or_tag(doc, mode, tag, image_path, width_in, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph(); p.alignment = align
    if mode == "template":
        r = p.add_run("{%" + tag + "}")
        r.font.size = Pt(10)
    else:
        p.add_run().add_picture(image_path, width=Inches(width_in))
    return p

# ------------------------------------------------------------------ build
def build(mode, out_path):
    T = mode == "template"
    def V(tag, sample):   # value: tag in template mode, real text in sample mode
        return "{" + tag + "}" if T else sample

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    st = doc.styles["Normal"]
    st.font.name = "Trebuchet MS"; st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Trebuchet MS")

    # ================= PAGE 1 — COVER =================
    t = doc.add_table(rows=1, cols=2); no_borders(t)
    c0, c1 = t.rows[0].cells
    c0.width = Inches(4.9); c1.width = Inches(2.1)
    fixed_widths(t, [4.9, 2.1])
    set_cell_bg(c0, YELLOW)
    c1p = c1.paragraphs[0]
    r = c1p.add_run("S.A.M.U.D.R.A"); r.font.bold = True; r.font.size = Pt(16); r.font.color.rgb = SLATE
    p = c1.add_paragraph(); r = p.add_run("PERFORMANCE.Intelligence")
    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = SLATE

    p = doc.add_paragraph()
    r = p.add_run("SAMUDRA"); r.font.size = Pt(34); r.font.color.rgb = SLATE
    r7 = p.add_run("7"); r7.font.size = Pt(15); r7.font.color.rgb = SLATE; r7.font.superscript = True
    r2 = p.add_run(".Jobfit"); r2.font.size = Pt(34); r2.font.color.rgb = SLATE

    big = doc.add_table(rows=1, cols=1); no_borders(big)
    fixed_widths(big, [7.0])
    cell = big.rows[0].cells[0]
    set_cell_bg(cell, YELLOW)
    tr = big.rows[0]._tr; trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight"); h.set(qn("w:val"), "10400"); h.set(qn("w:hRule"), "exact")
    trPr.append(h)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Report: "); r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK
    r = p.add_run("SAMUDRA.Jobfit  "); r.font.size = Pt(13); r.font.color.rgb = DARK
    r = p.add_run("ID: "); r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK
    r = p.add_run(V("report_ref", "SJF.1") + "  "); r.font.size = Pt(13); r.font.color.rgb = DARK
    r = p.add_run("Date: "); r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK
    r = p.add_run(V("report_date", "5th May 2026")); r.font.size = Pt(13); r.font.color.rgb = DARK
    for _ in range(13): cell.add_paragraph()
    pl = cell.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.add_run().add_picture(f"{ASSETS}/gps7_logo.png", width=Inches(1.7))

    para(doc, "@2025 GESTALT PSYCHO.GRAPHOLOGY INTELLIGENCE SYSTEM", size=9,
         color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
    para(doc, "Distributed by: Samudra7axis  II  graphoworldproject@gmail.com  II  "
              "www.graphoworld.com  II  +91 9095215551 ; +91 9994662290",
         size=8, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # ================= PAGE 2 — CANDIDATE + GPS7 + 7 TRAITS =================
    brand_header(doc)
    info = doc.add_table(rows=1, cols=3); no_borders(info)
    bar, left, right = info.rows[0].cells
    bar.width = Inches(0.15); left.width = Inches(3.3); right.width = Inches(3.55)
    fixed_widths(info, [0.15, 3.3, 3.55])
    set_cell_bg(bar, YELLOW)
    def info_line(cell, label, value):
        p = cell.add_paragraph() if cell.paragraphs[0].text or cell.paragraphs[0].runs else cell.paragraphs[0]
        r = p.add_run(label); r.font.size = Pt(12); r.font.color.rgb = STEEL
        r = p.add_run(value); r.font.size = Pt(12); r.font.color.rgb = STEEL
        p.paragraph_format.space_after = Pt(4)
    info_line(left,  "Date: ", V("report_date_short", "05.05.2026"))
    info_line(left,  "Name: ", V("candidate_name", "Priyanka K"))
    info_line(left,  "Education: ", V("education", "Media/Journalism"))
    info_line(left,  "Job Applied/Desired: ", V("job_applied", "Media"))
    info_line(right, "Assessment Input: ", V("assessment_input", "Handwriting + Signatures"))
    info_line(right, "Report Code: ", V("report_code", "SJF.MBA.001"))
    info_line(right, "Mobile: ", V("mobile", "+91 9994662290"))
    info_line(right, "Email Id: ", V("email", "creative1954@gmail.com"))

    section_title(doc, "Powered by GPS7")
    para(doc, "This report is powered by GPS7 – a proprietary Gestalt Grapho-Psychology "
              "Intelligence System and one of its kind in the world, designed to interpret "
              "handwriting and signatures as visible expressions of subconscious behavioural "
              "patterns. Using an integrated Gestalt Psycho.Graphology approach, the system "
              "evaluates rhythm, pressure, spacing, movement, form, and signature structure "
              "to create a unified psychological perspective rather than isolated trait "
              "readings. The framework draws upon established psychological research "
              "recognizing handwriting as a reflection of underlying cognitive and "
              "subconscious processes.", size=10.5, space_after=8)

    section_title(doc, "The Seven Core Performance Traits")
    fw = doc.add_table(rows=1, cols=2); no_borders(fw)
    b, img = fw.rows[0].cells
    b.width = Inches(0.15); img.width = Inches(6.85)
    fixed_widths(fw, [0.15, 6.85])
    set_cell_bg(b, YELLOW)
    ip = img.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(f"{ASSETS}/samudra_framework.png", width=Inches(5.9))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    defs = [
        ("Stability", "emotional balance, consistency, and dependability conditions."),
        ("Adaptability", "flexibility, openness to change, and the ability to adjust."),
        ("Motivation", "measures drive, initiative, ambition, and goal-orientation energy."),
        ("Unbiased", "objectivity, balanced judgment, balanced excessive emotional influence."),
        ("Decisioning", "clarity of thought, planning ability, and execution orientation."),
        ("Resilience", "recovery ability, stress tolerance, response to pressure or setbacks."),
        ("Assertiveness", "confidence, communication strength, and ability to take initiative."),
    ]
    for name, d in defs:
        p = doc.add_paragraph()
        r = p.add_run(f"{name} – "); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
        r = p.add_run(d); r.font.size = Pt(10); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.4)
    yellow_rule(doc)
    doc.add_page_break()

    # ================= PAGE 3 — BEHAVIORAL PROFILE (radar) + S A M U =================
    brand_header(doc)
    section_title(doc, "Behavioral Profile")
    image_or_tag(doc, mode, "radarChart", f"{ASSETS}/sample_radar.png", 5.1)

    trait_block(doc, "S – Social Intelligence (" + V("t_s_score", "68") + ")", V("t_s_text",
        "Comfortable interacting and collaborating within team environments. Builds "
        "functional relationships and adapts communication based on context. Handles "
        "routine interpersonal situations well, with scope to strengthen influence in "
        "complex or high-stake interactions."))
    trait_block(doc, "A – Adaptability (" + V("t_a_score", "82") + ")", V("t_a_text",
        "Highly flexible in approach, adjusting quickly to changing demands and "
        "environments. Responds well to dynamic situations and shifting priorities. "
        "Demonstrates openness to new ideas and ways of working, supporting performance "
        "in fast-paced roles."))
    trait_block(doc, "M – Mental Clarity (" + V("t_m_score", "74") + ")", V("t_m_text",
        "Demonstrates clear thinking and structured understanding in most situations. "
        "Processes information effectively and maintains focus on tasks. Able to "
        "interpret situations with reasonable clarity, supporting practical and timely "
        "responses."))
    trait_block(doc, "U – Unbiased (" + V("t_u_score", "60") + ")", V("t_u_text",
        "Shows a fair level of objectivity in decision-making. At times, judgment may be "
        "influenced by context or personal perspective. With conscious effort, can "
        "strengthen neutrality and improve balanced evaluation in critical situations."))
    yellow_rule(doc)
    doc.add_page_break()

    # ================= PAGE 4 — FIT SCORE STRIP + D R A =================
    brand_header(doc)
    section_title(doc, "Behavioral Profile")
    strip = doc.add_table(rows=1, cols=3); no_borders(strip)
    fixed_widths(strip, [2.33, 2.33, 2.33])
    labels = [("SAMUDRA FIT SCORE", V("samudra_fit_score", "71%"), "Strong Alignment"),
              ("STRONGEST TRAIT", V("strongest_trait", "A – Adaptability"), V("strongest_trait_score", "82")),
              ("FOCUS TRAIT", V("focus_trait", "U – Unbiased"), V("focus_trait_score", "60"))]
    for cell, (lab, big, sub) in zip(strip.rows[0].cells, labels):
        cell.width = Inches(2.33)
        set_cell_bg(cell, CREAM)
        cell_borders(cell, color="F5C542", size=6)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lab); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(big); r.font.bold = True; r.font.size = Pt(18); r.font.color.rgb = ORANGE
        p3 = cell.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p3.add_run(sub); r.font.size = Pt(10); r.font.color.rgb = ORANGE if lab != "SAMUDRA FIT SCORE" else DARK
        r.font.bold = lab != "SAMUDRA FIT SCORE"
    para(doc, "Scale: 0 (Lowest) to 100 (Highest)", size=9, bold=True, color=ORANGE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=12)

    trait_block(doc, "D – Drive & Ambition (" + V("t_d_score", "70") + ")", V("t_d_text",
        "Displays consistent intent to achieve goals with practical ambition. Maintains "
        "steady effort and engagement across responsibilities. Shows willingness to take "
        "initiative, especially in familiar or goal-oriented environments."))
    trait_block(doc, "R – Resilience (" + V("t_r_score", "66") + ")", V("t_r_text",
        "Handles routine pressures with reasonable stability and recovery ability. Can "
        "manage day-to-day stress effectively. Under prolonged or repetitive pressure, "
        "consistency may fluctuate, indicating scope to build stronger endurance."))
    trait_block(doc, "A – Accountability (" + V("t_ac_score", "78") + ")", V("t_ac_text",
        "Takes ownership of responsibilities with reliability and commitment. Follows "
        "through on tasks and demonstrates a responsible approach to work. Shows strong "
        "intent to deliver outcomes and uphold assigned roles."))
    yellow_rule(doc)
    doc.add_page_break()

    # ================= PAGE 5 — JOBFIT EVALUATION (gauge) =================
    brand_header(doc)
    section_title(doc, "Jobfit.Evaluation")
    image_or_tag(doc, mode, "gaugeChart", f"{ASSETS}/sample_gauge.png", 4.4)

    box = doc.add_table(rows=1, cols=1); no_borders(box)
    fixed_widths(box, [7.0])
    bc = box.rows[0].cells[0]
    set_cell_bg(bc, "EFF6FD"); cell_borders(bc, color="BBD4EE", size=6)
    p = bc.paragraphs[0]
    r = p.add_run("The overall Job Fit score of ")
    r.font.size = Pt(10.5); r.font.color.rgb = DARK
    r = p.add_run(V("jobfit_score", "78") + "%")
    r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = NAVY
    r = p.add_run(" " + V("jobfit_box_text",
        "reflects strong alignment with the job requirements. The candidate is "
        "well-suited to perform effectively in the applied role and has a higher "
        "likelihood of delivering consistent results."))
    r.font.size = Pt(10.5); r.font.color.rgb = DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for lab, col in [("  0% – 25%  Low Fit   ", YELLOW), ("  25% – 75%  Moderate to High Fit   ", "2FA84F"),
                     ("  75% – 100%  Very High Fit", "D9D9D9")]:
        sq = p.add_run("■ "); sq.font.color.rgb = RGBColor.from_string(col); sq.font.size = Pt(10)
        r = p.add_run(lab); r.font.size = Pt(9); r.font.color.rgb = DARK
    doc.add_paragraph()

    section_title(doc, "Jobfit.Summary")
    para(doc, V("jobfit_summary_p1",
        "The dial reflects the final Job Fit score of 78%, representing a strong "
        "alignment between the candidate's natural behavioral profile and the demands "
        "of the applied role. This score is derived from the SAMUDRA trait analysis "
        "(average ~71%) and is further refined through role alignment and industry fit "
        "mapping, ensuring that both inherent tendencies and external job requirements "
        "are considered in the evaluation."), size=10.5, space_after=8)
    para(doc, V("jobfit_summary_p2",
        "The needle indicates the candidate's position on the 0–100% fit scale, where "
        "higher placement signifies stronger compatibility with role expectations, work "
        "environment, and performance demands. The positioning within the higher range "
        "suggests that the individual is well-equipped to adapt, engage, and contribute "
        "effectively in relevant roles, with a higher likelihood of consistent "
        "performance and integration within professional settings."), size=10.5, space_after=8)
    yellow_rule(doc)
    doc.add_page_break()

    # ================= PAGE 6 — INDUSTRY MAPPING =================
    brand_header(doc)
    section_title(doc, "Industry Mapping")
    itbl = doc.add_table(rows=2 if T else 10, cols=3)
    itbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(itbl)
    fixed_widths(itbl, [2.7, 2.6, 1.7])
    widths = [Inches(2.7), Inches(2.6), Inches(1.7)]
    hdr = itbl.rows[0].cells
    for c, w, txt in zip(hdr, widths, ["Industry", "Fit %", "Fit Level"]):
        c.width = w; set_cell_bg(c, YELLOW)
        p = c.paragraphs[0]; r = p.add_run(txt)
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK

    sample_rows = [
        ("OTT / Streaming Content", 86, "HIGH"),
        ("Digital Media", 84, "HIGH"),
        ("Content Production", 82, "HIGH"),
        ("Television / Broadcasting", 78, "MODERATE–HIGH"),
        ("Advertising / Creative Agency", 74, "MODERATE"),
        ("Event & Live Production", 70, "MODERATE"),
        ("Corporate Communications", 66, "MODERATE–LOW"),
        ("Public Relations", 62, "LOW–MODERATE"),
        ("Data-driven Media Analytics", 54, "LOW"),
    ]
    def bar(pct):  # yellow block bar sized by pct
        full = round(pct / 100 * 14)
        return "▰" * full + "▱" * (14 - full)

    if T:
        row = itbl.rows[1].cells
        row[0].width = widths[0]; row[1].width = widths[1]; row[2].width = widths[2]
        set_cell_bg(row[0], CREAM); set_cell_bg(row[1], CREAM); set_cell_bg(row[2], CREAM)
        p = row[0].paragraphs[0]
        r = p.add_run("{#industries}{name}"); r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = STEEL
        p = row[1].paragraphs[0]
        r = p.add_run("{fitPct}%  "); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK
        r = p.add_run("{fitBar}"); r.font.size = Pt(9); r.font.color.rgb = ORANGE
        p = row[2].paragraphs[0]
        r = p.add_run("{fitLevel}{/industries}"); r.font.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(NAVYBADGE)
    else:
        for i, (name, pct, lvl) in enumerate(sample_rows):
            row = itbl.rows[i + 1].cells
            shade = CREAM if i < 3 else ("EFF4FA" if i < 6 else "F2F5F9")
            for c, w in zip(row, widths):
                c.width = w; set_cell_bg(c, shade)
            p = row[0].paragraphs[0]
            r = p.add_run(name); r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = STEEL
            p = row[1].paragraphs[0]
            r = p.add_run(f"{pct}%  "); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK
            r = p.add_run(bar(pct)); r.font.size = Pt(9); r.font.color.rgb = ORANGE
            p = row[2].paragraphs[0]
            r = p.add_run(lvl); r.font.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(NAVYBADGE) if pct >= 80 else STEEL
    doc.add_paragraph()

    section_title(doc, "Industry Fit")
    for tag, sample in [
        ("industry_fit_strong",
         "Strongest Industry Alignment: OTT / Streaming Content (86%), Digital Media (84%), "
         "and Content Production (82%) — indicating strong suitability for creative, "
         "fast-moving, and communication-driven environments."),
        ("industry_fit_moderate",
         "Moderate Alignment: Television / Broadcasting (78%), Advertising / Creative Agency "
         "(74%), and Event & Live Production (70%) — suggesting good potential with moderate "
         "adaptation to role demands and workflows."),
        ("industry_fit_development",
         "Development Areas: Corporate Communications (66%), Public Relations (62%), and "
         "Data-driven Media Analytics (54%) — where structured communication, consistency, "
         "or analytical orientation may require additional refinement."),
        ("industry_fit_overall",
         "Overall Observation: The profile demonstrates a stronger inclination toward "
         "creative, adaptive, and people-oriented roles rather than highly structured or "
         "data-intensive functions."),
    ]:
        para(doc, V(tag, sample), size=10, space_after=6)
    yellow_rule(doc)
    doc.add_page_break()

    # ================= PAGE 7 — STRENGTH / VARIABILITY ZONES =================
    brand_header(doc)
    section_title(doc, "Strength Zone")
    stbl = doc.add_table(rows=2, cols=4); no_borders(stbl)
    fixed_widths(stbl, [1.75, 1.75, 1.75, 1.75])
    s_defaults = [
        ("Comfortable working in dynamic environments",
         "Adapts easily to changing situations and fast-paced settings, contributing "
         "effectively to production and content workflows."),
        ("Strong communication ability",
         "Expresses ideas clearly and engages well with teams, essential for storytelling, "
         "coordination, and content execution."),
        ("Creative orientation",
         "Shows natural inclination towards content creation, ideas generation, and "
         "bringing originality to projects."),
        ("Positive team engagement",
         "Works well with others, maintains constructive interactions and supports team "
         "goals with a collaborative approach."),
    ]
    for i in range(4):
        top = stbl.rows[0].cells[i]; bot = stbl.rows[1].cells[i]
        top.width = Inches(1.75); bot.width = Inches(1.75)
        set_cell_bg(top, CREAM); set_cell_bg(bot, CREAM)
        cell_borders(top, color="F5C542", size=6, edges=("top", "left", "right"))
        cell_borders(bot, color="F5C542", size=6, edges=("left", "right", "bottom"))
        p = top.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(V(f"s{i+1}_title", s_defaults[i][0]))
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p = bot.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(V(f"s{i+1}_text", s_defaults[i][1]))
        r.font.size = Pt(8.5); r.font.color.rgb = DARK
    doc.add_paragraph()
    para(doc, V("strength_zone_summary",
        "These strengths collectively contribute to the individual's ability to perform "
        "effectively in creative, fast-paced, and collaborative environments. The profile "
        "reflects strong potential to adapt to evolving demands, manage responsibilities "
        "with flexibility, and deliver impactful outcomes consistently across "
        "communication-driven and execution-oriented roles."), size=10, space_after=8)

    section_title(doc, "Variability Zone")
    vtbl = doc.add_table(rows=2, cols=3); no_borders(vtbl)
    fixed_widths(vtbl, [2.33, 2.33, 2.33])
    v_defaults = [
        ("Decision-making may rely more on instinct",
         "Tends to take decisions based on experience and feel; may benefit from more "
         "structured analysis in complex situations."),
        ("Consistency may fluctuate under sustained pressure",
         "Performance can vary during prolonged or repetitive stress situations; requires "
         "breaks and reset for optimum output."),
        ("Objectivity may vary depending on context",
         "Personal perspective can influence judgment; conscious effort needed to maintain "
         "neutrality in critical evaluations."),
    ]
    for i in range(3):
        top = vtbl.rows[0].cells[i]; bot = vtbl.rows[1].cells[i]
        top.width = Inches(2.33); bot.width = Inches(2.33)
        set_cell_bg(top, CREAM); set_cell_bg(bot, CREAM)
        cell_borders(top, color="F5C542", size=6, edges=("top", "left", "right"))
        cell_borders(bot, color="F5C542", size=6, edges=("left", "right", "bottom"))
        p = top.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(V(f"v{i+1}_title", v_defaults[i][0]))
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p = bot.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(V(f"v{i+1}_text", v_defaults[i][1]))
        r.font.size = Pt(8.5); r.font.color.rgb = DARK
    doc.add_paragraph()
    para(doc, V("variability_zone_summary",
        "These variability areas are not limitations but natural behavioral tendencies "
        "that can be effectively managed through awareness, structured thinking, and the "
        "right professional strategies. Proactive planning, disciplined execution, and "
        "mindful self-regulation can help maintain balance, improve consistency, and "
        "support stronger decision-making across demanding work environments."),
        size=9.5, space_after=6)
    yellow_rule(doc)
    doc.add_page_break()

    # ================= OVERLAYS (O-Series, template loop) =================
    if T:
        p = doc.add_paragraph()
        r = p.add_run("{#hasOverlays}"); r.font.size = Pt(2)
        section_title(doc, "Applied Overlays")
        p = doc.add_paragraph()
        r = p.add_run("{#overlays}"); r.font.size = Pt(2)
        p = doc.add_paragraph()
        r = p.add_run("{code} – {title}")
        r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = STEEL
        p = doc.add_paragraph()
        r = p.add_run("{body}"); r.font.size = Pt(11); r.font.color.rgb = BODY
        p = doc.add_paragraph()
        r = p.add_run("{/overlays}"); r.font.size = Pt(2)
        # The page break lives INSIDE the conditional: when no overlays are
        # enabled the whole block is removed and no blank page is left behind.
        p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
        r = p.add_run("{/hasOverlays}"); r.font.size = Pt(2)

    # ================= PAGE 8 — EXECUTIVE SUMMARY + DISCLAIMER =================
    brand_header(doc)
    section_title(doc, "Executive Summary")
    es = doc.add_table(rows=2, cols=4); no_borders(es)
    fixed_widths(es, [1.75, 1.75, 1.75, 1.75])
    heads = ["Job Fit", "Top Strength", "Key Variability", "Best Industries"]
    vals = [(V("jobfit_score", "78") + "%", True),
            (V("top_strength", "Adaptability & Communication"), False),
            (V("key_variability", "Structured Decision-Making"), False),
            (V("best_industries", "OTT, Digital Media, Content Production"), False)]
    for i in range(4):
        top = es.rows[0].cells[i]; bot = es.rows[1].cells[i]
        top.width = Inches(1.75); bot.width = Inches(1.75)
        cell_borders(top, color="F5C542", size=6, edges=("top", "left", "right"))
        cell_borders(bot, color="F5C542", size=6, edges=("left", "right", "bottom"))
        p = top.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(heads[i]); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
        p = bot.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        txt, big = vals[i]
        r = p.add_run(txt); r.font.bold = True
        r.font.size = Pt(18) if big else Pt(9.5)
        r.font.color.rgb = ORANGE
    doc.add_paragraph()

    exe = doc.add_table(rows=1, cols=2); no_borders(exe)
    ebar, ebody = exe.rows[0].cells
    ebar.width = Inches(0.15); ebody.width = Inches(6.85)
    fixed_widths(exe, [0.15, 6.85])
    set_cell_bg(ebar, YELLOW)
    for i, (tag, sample) in enumerate([
        ("exec_summary_p1",
         "The individual demonstrates a strong combination of adaptability and "
         "communication, enabling effective interaction, coordination, and responsiveness "
         "across changing work environments. The profile reflects an ability to adjust "
         "quickly to evolving demands while maintaining clarity in communication and "
         "collaboration with teams and stakeholders."),
        ("exec_summary_p2",
         "The assessment also highlights a structured decision-making style, suggesting a "
         "preference for organized planning, logical evaluation, and disciplined execution. "
         "Such individuals generally perform well in environments that require balancing "
         "creativity with coordination and operational consistency."),
        ("exec_summary_p3",
         "Industry alignment indicates stronger suitability for sectors such as OTT "
         "platforms, digital media, and content production, where communication flow, "
         "content coordination, deadline management, and execution efficiency play an "
         "important role in day-to-day performance and long-term professional growth."),
    ]):
        p = ebody.paragraphs[0] if i == 0 else ebody.add_paragraph()
        r = p.add_run(V(tag, sample)); r.font.size = Pt(10.5); r.font.color.rgb = BODY
        p.paragraph_format.space_after = Pt(8)

    pass
    dis = doc.add_table(rows=1, cols=2); no_borders(dis)
    dbar, dbody = dis.rows[0].cells
    dbar.width = Inches(0.15); dbody.width = Inches(6.85)
    fixed_widths(dis, [0.15, 6.85])
    set_cell_bg(dbar, YELLOW)
    p = dbody.paragraphs[0]
    r = p.add_run("Disclaimer"); r.font.bold = True; r.font.size = Pt(16); r.font.color.rgb = STEEL
    p = dbody.add_paragraph()
    r = p.add_run(
        "This report has been generated with the knowledge and consent of the applicant "
        "and is based on the analysis of provided handwriting and signature samples. The "
        "interpretation reflects behavioral tendencies derived through the GPS7 Gestalt "
        "Grapho.Psychology framework and is intended to support career awareness and "
        "guidance. It should be used as an indicative tool alongside other evaluation "
        "methods and not as the sole basis for decision-making.")
    r.font.size = Pt(10); r.font.color.rgb = BODY
    para(doc, "@2025 GESTALT PSYCHO.GRAPHOLOGY INTELLIGENCE SYSTEM", size=9, color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    para(doc, "Distributed by: Samudra7axis  II  graphoworldproject@gmail.com  II  "
              "www.graphoworld.com  II  +91 9095215551 ; +91 9994662290  II  END OF THE REPORT",
         size=8, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.save(out_path)
    print(f"built [{mode}] -> {out_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
